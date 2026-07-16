"""python-docx DOCX chunk highlighter.

Strategy
--------
* ``python-docx`` iterates every paragraph (body + table cells).
* For each retrieved snippet the full snippet token sequence is checked
  against each paragraph using a bidirectional containment test:

  1. *Paragraph-in-snippet*: the paragraph's tokens appear as a contiguous
     sub-sequence inside the snippet tokens.  This matches paragraphs that
     were merged together to form the chunk text (common with HeadingChunker
     breadcrumb mode).
  2. *Fragment-in-paragraph*: a significant line/sentence fragment from the
     snippet appears in the paragraph tokens.  This is the fallback for
     long or split paragraphs.

* Short paragraphs (fewer than ``_MIN_PARA_TOKENS`` tokens) are only
  considered for direction 2 to avoid spurious matches on titles/labels.
* Highlighting is applied by injecting a ``<w:highlight w:val="..."/>``
  element into each matching run's ``<w:rPr>`` via the python-docx XML API.
* The modified document is serialised in-memory; the source file is never
  touched.
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Sequence

from VisualMarkers.VisualMarker import ChunkSnippet, VisualMarker

# Minimum number of tokens a paragraph must have to be considered for the
# paragraph-in-snippet (direction 1) test.  Short tokens like "Note:" or
# "Introduction" would otherwise match almost anything.
_MIN_PARA_TOKENS = 4

# Minimum character length of a snippet line/sentence fragment to bother
# searching for in direction 2.
_MIN_FRAGMENT_LEN = 12

_PUNCT_RX = re.compile(r"[^\w\s]")

# WD_COLOR_INDEX name → w:highlight val attribute value (lowercased strings
# accepted by the OOXML spec for w:highlight/@w:val).
_WD_COLORS: dict[str, str] = {
    "yellow": "yellow",
    "green": "green",
    "cyan": "cyan",
    "magenta": "darkMagenta",
    "pink": "pink",
    "red": "red",
    "blue": "blue",
    "darkblue": "darkBlue",
    "orange": "yellow",  # no direct orange in WD_COLOR_INDEX
    "lime": "green",
}


def _parse_color(color: str) -> str:
    """Map a CSS color name or ``#RRGGBB`` string to a ``w:highlight`` value.

    Falls back to ``"yellow"`` for unknown inputs.
    """
    c = color.strip().lower().lstrip("#")
    # Try named-color lookup (strip any leading # first)
    named = color.strip().lower()
    if named in _WD_COLORS:
        return _WD_COLORS[named]
    return "yellow"


class DocxVisualMarker(VisualMarker):
    """Highlight retrieved chunks in a DOCX file using python-docx."""

    def mark_to_bytes(
        self,
        source_path: Path,
        snippets: Sequence[ChunkSnippet],
        *,
        highlight_color: str = "yellow",
    ) -> bytes:
        """Return *source_path* with *snippets* highlighted, as DOCX bytes."""
        try:
            from docx import Document  # type: ignore[import-not-found]
            from docx.oxml import OxmlElement  # type: ignore[import-not-found]
            from docx.oxml.ns import qn  # type: ignore[import-not-found]
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX visual marking. "
                "Install with: pip install python-docx"
            ) from exc

        wd_color = _parse_color(highlight_color)
        doc = Document(str(source_path))

        for snippet in snippets:
            text = (snippet.text or "").strip()
            if not text:
                continue
            snippet_tokens = _tokenize(text)
            if not snippet_tokens:
                continue

            # Collect all paragraphs: body + table cells
            all_paras = list(doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_paras.extend(cell.paragraphs)

            for para in all_paras:
                para_text = para.text.strip()
                if not para_text:
                    continue
                para_tokens = _tokenize(para_text)
                if _para_matches(para_tokens, snippet_tokens, text):
                    _highlight_paragraph(para, wd_color, qn, OxmlElement)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


# ---------------------------------------------------------------------------
# Matching helpers
# ---------------------------------------------------------------------------


def _para_matches(
    para_tokens: list[str],
    snippet_tokens: list[str],
    snippet_text: str,
) -> bool:
    """Return True if *para_tokens* is relevant to *snippet_tokens*."""
    # Direction 1: paragraph text is a contiguous sub-sequence of the snippet.
    # Only apply to non-trivial paragraphs to avoid false positives on labels.
    if len(para_tokens) >= _MIN_PARA_TOKENS and _contains_sequence(
        snippet_tokens, para_tokens
    ):
        return True

    # Direction 2 (fallback): a significant fragment of the snippet text
    # appears in the paragraph.
    for frag in _iter_fragments(snippet_text):
        frag_tokens = _tokenize(frag)
        if frag_tokens and _contains_sequence(para_tokens, frag_tokens):
            return True

    return False


def _contains_sequence(haystack: list[str], needle: list[str]) -> bool:
    """Return True if *needle* appears as a contiguous sub-sequence in *haystack*."""
    n = len(needle)
    if n == 0 or len(haystack) < n:
        return False
    for i in range(len(haystack) - n + 1):
        if haystack[i : i + n] == needle:
            return True
    return False


def _tokenize(text: str) -> list[str]:
    """Lower-case, strip punctuation, split into non-empty tokens."""
    return _PUNCT_RX.sub(" ", text.lower()).split()


def _iter_fragments(text: str):
    """Yield non-trivial line/sentence substrings of *text*."""
    seen: set[str] = set()
    for line in text.splitlines():
        stripped = line.strip()
        if len(stripped) >= _MIN_FRAGMENT_LEN and stripped not in seen:
            seen.add(stripped)
            yield stripped
    for sent in re.split(r"(?<=[\.\!\?])\s+", text):
        stripped = sent.strip()
        if len(stripped) >= _MIN_FRAGMENT_LEN and stripped not in seen:
            seen.add(stripped)
            yield stripped


# ---------------------------------------------------------------------------
# DOCX XML helpers
# ---------------------------------------------------------------------------


def _highlight_paragraph(para, wd_color: str, qn, OxmlElement) -> None:
    """Apply highlight to every run in *para*."""
    for run in para.runs:
        rPr = run._r.get_or_add_rPr()
        # Remove any existing highlight element.
        for existing in rPr.findall(qn("w:highlight")):
            rPr.remove(existing)
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), wd_color)
        rPr.append(hl)
